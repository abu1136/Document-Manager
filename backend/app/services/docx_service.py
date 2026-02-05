from docx import Document
import os

def generate_docx(title, content, document_id, output_path):
    """
    Generate a DOCX file with the given content
    """
    # Ensure output directory exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    d = Document()
    d.add_heading(title, 1)
    d.add_paragraph(content)
    d.add_paragraph("")
    d.add_paragraph(f"Document ID: {document_id}", style='Caption')
    d.save(output_path)
